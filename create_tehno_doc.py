#!/usr/bin/env python3
"""
Документ с точными данными из doc.txt и second.txt
Total pozitii 49, Suma 132'906.000 MDL
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill_color)
    tcPr.append(shading)

def set_repeat_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)

def create_document():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)

    # === ШАПКА ===
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = True

    left_cell = header_table.cell(0, 0)
    p = left_cell.paragraphs[0]
    run = p.add_run("TEHNO")
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = 'Arial Black'

    p2 = left_cell.add_paragraph()
    run2 = p2.add_run("Ofertă Nr 140126-03169 din 14.01.26")
    run2.bold = True
    run2.font.size = Pt(11)

    p3 = left_cell.add_paragraph()
    run3 = p3.add_run("Tel: 062163090; https://atehno.md")
    run3.font.size = Pt(10)

    right_cell = header_table.cell(0, 1)
    pr = right_cell.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_r = pr.add_run("Cumpărător:")
    run_r.bold = True
    run_r.font.size = Pt(11)

    pr2 = right_cell.add_paragraph()
    pr2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_r2 = pr2.add_run('ACADEMIA "STEFAN CEL MARE" A MAI')
    run_r2.bold = True
    run_r2.font.size = Pt(11)

    doc.add_paragraph()

    # === ТОЧНЫЕ ДАННЫЕ ИЗ doc.txt и second.txt ===
    items = [
        # doc.txt позиции
        ("OR-H285A/CRG725", "ORINK OR-H285A/CRG725 HP LJ P1100/1102/M1130/1132/1136/1210/1212/1214/1217;Canon LBP6000/6018/6020/6030/MF3010 (1.600p)", 70, "buc", 150.000, 10500.000),
        ("OK-Q2612AFX10", "Oklnk OK-Q2612AFX10/CRG-703 HP LJ 1010/1012/1015/1018/1020/1022/3015/3020/3030/3050/3052/3055/M1005/1319/Canon LBP2900/3000MF4010/4018/4110/4120/4122/4140/4150/4270/4320/4330/4340/4350/4370/4380/4660/4680/4690/L75/100/120/140/160 (2.000p)", 6, "buc", 150.000, 900.000),
        ("OK-CE278A/CRG-728", "Oklnk OK-CE278A/CRG-728 HP LJ P1566/1606/M1536/Canon 4570/4580/4730/4750/4780/4870/4890/L150/170/D520/550 LBP6200/6230/MF4410/4412/4420/4430/4450/4452/4550/4552 (2.000p)", 2, "buc", 155.000, 310.000),
        ("OK-CF230A/CRG-051", "Oklnk OK-CF230A/CRG-051 HP LJ Pro M203/227; Canon LBP162DW/164DW/MF264W/267DW/269DW (1.600p)", 4, "buc", 155.000, 620.000),
        ("OK-CF226X/CRG-052 H", "Oklnk OK-CF226X/CRG-052H HP LJ Pro M402/426, Canon LBP212/214/215/MF421/424/426/428/429 (9.200p)", 6, "buc", 145.000, 870.000),
        ("C-EXV42", "Toner Canon C-EXV42 Black (486g/appr. 10 200 pages 6%) for iR2206.2206N,2204.2204N,2204F,2202,2202N,2202i", 6, "buc", 225.000, 1350.000),
        ("OK-CF259X/CRG-057 H", "Oklnk OK-CF259X/CRG-057H HP LJ Pro M404/MFP M428; Canon LBP223/226/228/233/236/MF433/443/445/446/449/453/455/X1238, CRG-T08, w/o chip (10.000p)", 3, "buc", 755.000, 2265.000),
        ("OK-CF259X/CRG-057 H", "Oklnk OK-CF259X/CRG-057H HP LJ Pro M404/MFP M428; Canon LBP223/226/228/233/236/MF433/443/445/446/449/453/455/X1238, CRG-T08, w/o chip (10.000p)", 4, "buc", 235.000, 940.000),
        ("OK-CF259X/CRG-057 H", "Oklnk OK-CF259X/CRG-057H HP LJ Pro M404/MFP M428; Canon LBP223/226/228/233/236/MF433/443/445/446/449/453/455/X1238, CRG-T08, w/o chip (10.000p)", 10, "buc", 235.000, 2350.000),
        ("OK-CF217A/CRG-047", "Oklnk OK-CF217A/CRG-047 HP LJ Pro M102/104/MFP M129/130/132; Canon LBP112/113W/MF112/113 (1.600p)", 2, "buc", 140.000, 280.000),
        ("OK-CF219A/CRG-049", "OkInk OK-CF219A/CRG-049 Drum Unit HP LJ Pro M102/104/MFP M130/132; Canon LBP112/113W/MF112/113 (12.000p)", 2, "buc", 170.000, 340.000),
        ("OR-H285A/CRG725", "ORINK OR-H285A/CRG725 HP LJ P1100/1102/M1130/1132/1136/1210/1212/1214/1217;Canon LBP6000/6018/6020/6030/MF3010 (1.600p)", 2, "buc", 145.000, 290.000),
        ("OK-Q2612AFX10", "Oklnk OK-Q2612AFX10/CRG-703 HP LJ 1010/1012/1015/1018/1020/1022/3015/3020/3030/3050/3052/3055/M1005/1319/Canon LBP2900/3000MF4010/4018/4110/4120/4122/4140/4150/4270/4320/4330/4340/4350/4370/4380/4660/4680/4690/L75/100/120/140/160 (2.000p)", 2, "buc", 155.000, 310.000),
        ("OK-CE505X/CRG719 H", "Oklnk OK-CE505X/CF280X/CRG-719H HP LJ P2050/2055/Pro M401/425/Canon LBP251/252/253/6300/6310/6650/6670/6680/MF411/416/418/419/5840/5850/5870/5880/5930/5940/5950/5960/5980/6140/6160/6180/iR1133 (6.500p)", 2, "buc", 225.000, 450.000),
        ("OR-BTN2590", "ORINK OR-BTN2590 Brother HL-2402/2400/DCP-2860/2640/MFC-2802", 6, "buc", 245.000, 1470.000),
        ("OR-BDR2590", "ORINK OR-BDR2590 Brother HL-2402/2400/DCP-2860/2640/MFC-2802", 2, "buc", 245.000, 490.000),
        ("OK-W1350A/CRG-071", "Oklnk OK-W1350A/CRG-071 HP LJ M209/MFP M233/234, Canon LBP122/MF272dw/MF275dw, w/o chip (1.000p)", 3, "buc", 170.000, 510.000),
        ("CCT-CRG046HB/CF410X", "Starlnk CCT-CRG-046HB/CF410X Black Canon LBP653/654/MF732/734/735; HP CLJ Pro M377/452/477 (6.300p)", 1, "buc", 240.000, 240.000),
        ("CCT-CRG046HC/CF411X", "Starlnk CCT-CRG-046HC/CF411X Cyan Canon LBP653/654/MF732/734/735; HP CLJ Pro M377/452/477 (5.000p)", 1, "buc", 240.000, 240.000),
        ("CCT-CRG046HM/CF413X", "Starlnk CCT-CRG-046HM/CF413X Magenta Canon LBP653/654/MF732/734/735; HP CLJ Pro M377/452/477 (5.000p)", 1, "buc", 240.000, 240.000),
        ("CCT-CRG046HY/CF412X", "Starlnk CCT-CRG-046HY/CF412X Yellow Canon LBP653/654/MF732/734/735; HP CLJ Pro M377/452/477 (5.000p)", 1, "buc", 240.000, 240.000),
        ("9857901857032", "Ink Cartridge Epson T55K60N UltraChrome HDX/HD 700ml, Vivid Light Magenta, C13T55K60N For Epson SC-P6000/7000/9000", 2, "buc", 5149.000, 10298.000),
        ("9857901857032", "Ink Cartridge Epson T55K200 UltraChrome HDX/HD 700ml, Cyan / C13T804200 For Epson SC-P6000/7000/9000", 2, "buc", 5149.000, 10298.000),
        # second.txt позиции
        ("9857901644113", "Ink Cartridge Epson T55K400 UltraChrome HDX/HD 700ml, Yellow / C13T804400 For Epson SC-P6000/7000/9000", 2, "buc", 5149.000, 10298.000),
        ("9857901551510", "Epson SureColor SC-P6000 STD/ SC-P7000 STD/ SC-P8000 STD/ SC-P9000 STD Ink", 2, "buc", 5149.000, 10298.000),
        ("9857901644168", "Ink Cartridge Epson T55KA00 UltraChrome HDX/HD 700ml, Orange /C13T804A00 For Epson SC-P6000/7000/9000", 2, "buc", 5149.000, 10298.000),
        ("9857901644175", "Ink Cartridge Epson T55KB00 UltraChrome HDX/HD 700ml, Green / C13T804B00 For Epson SC-P6000/7000/9000", 2, "buc", 5149.000, 10298.000),
        ("9857901644144", "Ink Cartridge Epson T55K700 UltraChrome HDX/HD 700ml, Light Black / C13T804700 For Epson SC-P6000/7000/9000", 2, "buc", 5149.000, 10298.000),
        ("9857902614979", "415X 415A Multipack with Chip Compatible for HP 415X W2030X Toner Cartridge for HP Color Laserjet Pro MFP M479fdw M454dw M479fnw M479dw M479fdn M454dn Black Cyan Yellow Magenta", 1, "buc", 6000.000, 6000.000),
        ("CCT-CRG046HC/CF411X", "Starlnk CCT-CRG-046HC/CF411X Cyan Canon LBP653/654/MF732/734/735; HP CLJ Pro M377/452/477 (5.000p)", 1, "buc", 240.000, 240.000),
        ("CCT-CRG046HM/CF413X", "Starlnk CCT-CRG-046HM/CF413X Magenta Canon LBP653/654/MF732/734/735; HP CLJ Pro M377/452/477 (5.000p)", 1, "buc", 240.000, 240.000),
        ("CCT-CRG046HY/CF412X", "Starlnk CCT-CRG-046HY/CF412X Yellow Canon LBP653/654/MF732/734/735; HP CLJ Pro M377/452/477 (5.000p)", 1, "buc", 240.000, 240.000),
        ("OR-TN321K", "ORINK OR-TN321K Black Toner Tube Konica Minolta Bizhub C224/284/364, A33K150 (27.000p)", 1, "buc", 240.000, 240.000),
        ("OR-TN321C", "ORINK OR-TN321C Cyan Toner Tube Konica Minolta Bizhub C224/284/364, A33K450 (25.000p)", 1, "buc", 600.000, 600.000),
        ("OR-TN321Y", "ORINK OR-TN321Y Yellow Toner Tube Konica Minolta Bizhub C224/284/364, A33K250 (25.000p)", 1, "buc", 800.000, 800.000),
        ("CBP-TK1150", "Starlnk CBP-TK1150 Kyocera Ecosys P2235/M2135/2635/2735 (3.000p)", 1, "buc", 800.000, 800.000),
        ("CBT-106R01410", "Starink CBT-106R01410 Xerox WorkCentre 4250/4260 (25.000p)", 1, "buc", 120.000, 120.000),
        ("9359844620075", "Toner Sharp MX-B42T, Black Toner cartridge; 20k, for Sharp MX-B427WEU", 1, "buc", 1100.000, 1100.000),
        ("CBP-TK1150", "Starlnk CBP-TK1150 Kyocera Ecosys P2235/M2135/2635/2735 (3.000p)", 1, "buc", 3500.000, 3500.000),
        ("9857902615044", "Cartus Kyocera TK-5240", 1, "buc", 120.000, 120.000),
        ("9857902615051", "Cartus Kyocera TK-5240 Black", 1, "buc", 715.000, 715.000),
        ("9857902615068", "Cartus Kyocera TK-5240 Cyan", 1, "buc", 700.000, 700.000),
        ("9857902615068", "Cartus Kyocera TK-5240 Magenta", 1, "buc", 700.000, 700.000),
        ("9857902615075", "Cartus Kyocera TK-5240 Yellow", 1, "buc", 700.000, 700.000),
        ("9857902614924", "Cartus toner compatibil Ricoh PRO C5100 800g Black OEM", 1, "buc", 2000.000, 2000.000),
        ("9857902614931", "Cartus toner compatibil Ricoh PRO C5100 800g Cyan OEM", 1, "buc", 2000.000, 2000.000),
        ("9857902614948", "Cartus toner compatibil Ricoh PRO C5100 800g Magenta OEM", 1, "buc", 2000.000, 2000.000),
        ("9857902614955", "Cartus toner compatibil Ricoh PRO C5100 800g Yellow OEM", 1, "buc", 2000.000, 2000.000),
        ("9857902614962", "C8000 Toner Cartridge Compatible for Xerox Versalink C8000 C8000DT C8000W C8000DTM Printers, High Yield 12600 Pages, with Chip (Multicolor)", 1, "buc", 10800.000, 10800.000),
    ]

    # Создание таблицы
    table = doc.add_table(rows=len(items) + 1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Cm(3.5), Cm(13.5), Cm(1.5), Cm(1.5), Cm(2.5), Cm(3)]
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width

    headers = ['Articol', 'Bunuri', 'Cant.', 'Un.', 'Preț', 'Suma']
    header_row = table.rows[0]
    set_repeat_header(header_row)

    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, 'D9D9D9')

    for row_idx, item in enumerate(items):
        row = table.rows[row_idx + 1]
        for col_idx, value in enumerate(item):
            cell = row.cells[col_idx]

            if col_idx == 4:
                cell.text = f"{value:,.3f}".replace(',', "'").replace('.', ',')
            elif col_idx == 5:
                cell.text = f"{value:,.3f}".replace(',', "'").replace('.', ',')
            else:
                cell.text = str(value)

            for para in cell.paragraphs:
                if col_idx in [0, 2, 3]:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif col_idx in [4, 5]:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in para.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph()

    total_items = len(items)
    total_suma = sum(item[5] for item in items)

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run1 = summary.add_run(f"Total poziții: {total_items}\n")
    run1.bold = True
    run1.font.size = Pt(11)

    suma_str = f"{total_suma:,.3f}".replace(',', "'").replace('.', ',')
    run2 = summary.add_run(f"Suma: {suma_str} MDL\n")
    run2.bold = True
    run2.font.size = Pt(11)

    tva_amount = 22150.980
    tva_str = f"{tva_amount:,.3f}".replace(',', "'").replace('.', ',')
    run3 = summary.add_run(f"Inclusiv TVA: {tva_str}\n")
    run3.bold = True
    run3.font.size = Pt(11)

    run4 = summary.add_run("Una suta treizeci doua mii noua sute sase lei 00 bani")
    run4.italic = True
    run4.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    sign_table = doc.add_table(rows=1, cols=2)
    sign_table.autofit = True

    left_sign = sign_table.cell(0, 0)
    p_left = left_sign.paragraphs[0]
    run_left = p_left.add_run("Eliberat: _______________ Badalov Alexei")
    run_left.font.size = Pt(11)

    right_sign = sign_table.cell(0, 1)
    p_right = right_sign.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_right = p_right.add_run("Primit: _______________")
    run_right.font.size = Pt(11)

    doc.save('TEHNO_document.docx')
    print("✓ Документ создан: TEHNO_document.docx")
    print(f"  Total poziții: {total_items}")
    print(f"  Suma: {suma_str} MDL")

if __name__ == "__main__":
    create_document()
